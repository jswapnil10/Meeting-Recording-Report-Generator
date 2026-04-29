import json
import math
import os
import re
import subprocess
import sys
import tempfile
import time

import boto3
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from google import genai
from google.genai import types
from google.oauth2 import service_account


def get_gemini_client():
    session = boto3.Session(region_name='us-east-1')
    sm = session.client('secretsmanager')

    gemini_secret_values = json.loads(
        sm.get_secret_value(SecretId="bos/google/gemini")['SecretString']
    )
    private_key_str = gemini_secret_values["private_key"]
    gemini_secret_values["private_key"] = private_key_str.replace("\\n", "\n")

    scopes = ["https://www.googleapis.com/auth/cloud-platform"]
    credentials = service_account.Credentials.from_service_account_info(
        gemini_secret_values, scopes=scopes
    )

    client = genai.Client(
        credentials=credentials,
        vertexai=True,
        project=gemini_secret_values.get('project_id'),
        location="global",
    )
    return client


def get_video_duration_seconds(video_path):
    result = subprocess.run(
        ["ffprobe", "-v", "error", "-show_entries", "format=duration",
         "-of", "default=noprint_wrappers=1:nokey=1", video_path],
        capture_output=True, text=True
    )
    return float(result.stdout.strip())


def split_video_into_chunks(video_path, chunk_duration_sec=1800, work_dir=None):

    if work_dir is None:
        work_dir = tempfile.mkdtemp(prefix="kt_chunks_")

    total_duration = get_video_duration_seconds(video_path)
    num_chunks = math.ceil(total_duration / chunk_duration_sec)
    chunk_paths = []

    print(f"Video duration: {total_duration:.0f}s ({total_duration/60:.1f} min)")
    print(f"Splitting into {num_chunks} chunk(s) of up to {chunk_duration_sec//60} min each...")

    for i in range(num_chunks):
        start = i * chunk_duration_sec
        chunk_path = os.path.join(work_dir, f"chunk_{i:03d}.mp4")
        cmd = [
            "ffmpeg", "-y", "-i", video_path,
            "-ss", str(start), "-t", str(chunk_duration_sec),
            "-c", "copy",
            # "-an",
            chunk_path
        ]
        subprocess.run(cmd, capture_output=True, check=True)
        chunk_paths.append(chunk_path)
        print(f"  Created chunk {i+1}/{num_chunks}: {chunk_path}")

    return chunk_paths, total_duration


def extract_screenshot(video_path, timestamp_sec, output_path):

    output_path = os.path.splitext(output_path)[0] + '.png'
    cmd = [
        "ffmpeg", "-y",
        "-i", video_path,
        "-ss", str(timestamp_sec),
        "-frames:v", "1",
        "-update", "1",
        output_path
    ]
    result = subprocess.run(cmd, capture_output=True)
    if result.returncode != 0 or not os.path.exists(output_path) or os.path.getsize(output_path) < 100:
        return None
    return output_path


def load_transcript(transcript_path):
    if not transcript_path or not os.path.exists(transcript_path):
        return None
    if transcript_path.lower().endswith('.docx'):
        doc = Document(transcript_path)
        return '\n'.join(p.text for p in doc.paragraphs)
    with open(transcript_path, "r", encoding="utf-8") as f:
        return f.read()


def get_transcript_segment(full_transcript, start_sec, end_sec):

    if not full_transcript:
        return None

    ts_pattern = re.compile(r'[\[\(]?(\d{1,2}):(\d{2}):(\d{2})[\]\)]?')
    lines = full_transcript.split('\n')
    segment_lines = []
    found_timestamps = False

    for line in lines:
        match = ts_pattern.search(line)
        if match:
            found_timestamps = True
            h, m, s = int(match.group(1)), int(match.group(2)), int(match.group(3))
            line_sec = h * 3600 + m * 60 + s
            if start_sec <= line_sec < end_sec:
                segment_lines.append(line)
        elif not found_timestamps:
            segment_lines.append(line)
        elif segment_lines:
            segment_lines.append(line)

    if found_timestamps and segment_lines:
        return '\n'.join(segment_lines)

    if start_sec == 0:
        return full_transcript
    return None


ANALYSIS_PROMPT = """You are an expert technical documentation writer. You are analyzing a segment of a Knowledge Transfer (KT) / team meeting recording.

This is segment {chunk_number} of {total_chunks} of the full meeting.

Your task is to produce an EXHAUSTIVE and DETAILED summary of everything discussed and shown in this video segment.

**You MUST cover:**

1. **Topics & Agenda Items**: Every distinct topic, feature, system, or concept discussed. Use clear headings.

2. **Technical Details**:
   - Architecture, system design, data flows explained
   - Code walkthroughs — describe what code/files were shown, key logic, patterns
   - APIs, endpoints, configurations, environment details
   - Database schemas, queries, data models
   - Tools, frameworks, libraries mentioned and how they are used

3. **Demonstrations & Walkthroughs**:
   - Step-by-step description of any demos shown
   - UI screens, dashboards, or tools navigated — describe what was on screen
   - Commands executed in terminal, their outputs

4. **Key Decisions & Action Items**:
   - Any decisions made during the meeting
   - Action items, owners, deadlines mentioned
   - Open questions or parking-lot items

5. **Important Visuals / Screenshots**:
   - Identify 3-8 moments in this video segment where the screen shows something important (architecture diagram, code, UI, config, terminal output, etc.)
   - For EACH, provide the EXACT timestamp in the format `[SCREENSHOT:MM:SS]` and a description of what is shown.
   - Choose moments that would be most valuable for someone reading the document later.

6. **Q&A and Discussion Points**:
   - Questions asked by team members and the answers given
   - Any confusion points that were clarified

**Formatting rules:**
- Use Markdown formatting (headings ##, ###, bullet points, bold, code blocks)
- Be verbose and thorough — this document will be the ONLY reference for people who missed the meeting
- Include actual content/values shown, not just "they showed a config file" — describe WHAT was in it
- For code shown on screen, include the key snippets if you can read them
- Timestamps for screenshots must be relative to THIS video segment (starting from 00:00)
- CRITICAL: Do NOT start with any intro/preamble like "Here is the documentation..." or "This segment covers..." — your VERY FIRST LINE must be a markdown heading (## or ###). Jump straight into content.
- This segment's content will be merged with other segments into one document, so write it as a continuation — do NOT repeat topics already covered in previous segments

{previous_context}

{transcript_section}

Now analyze the video thoroughly and produce the exhaustive documentation."""


def analyze_chunk(client, model, chunk_path, chunk_index, total_chunks, chunk_start_sec, transcript_segment, previous_summary=None):

    print(f"\n--- Analyzing Chunk {chunk_index + 1}/{total_chunks} (starts at {chunk_start_sec//60:.0f}m) ---")

    file_size_mb = os.path.getsize(chunk_path) / (1024 * 1024)
    print(f"  Sending {os.path.basename(chunk_path)} ({file_size_mb:.1f} MB) inline to Gemini...")

    with open(chunk_path, "rb") as f:
        video_bytes = f.read()

    transcript_section = ""
    if transcript_segment:
        transcript_section = f"""
**Meeting transcript for this segment is provided below for additional context. Use it alongside the video to ensure nothing is missed:**

```
{transcript_segment}
```
"""

    previous_context = ""
    if previous_summary:

        trimmed = previous_summary[:3000]
        previous_context = f"""
**Here is a summary of what was already covered in the PREVIOUS segment(s). Do NOT repeat this — continue from where it left off:**

```
{trimmed}
```
"""

    prompt = ANALYSIS_PROMPT.format(
        chunk_number=chunk_index + 1,
        total_chunks=total_chunks,
        transcript_section=transcript_section,
        previous_context=previous_context,
    )

    response = client.models.generate_content(
        model=model,
        contents=[
            types.Content(
                role="user",
                parts=[
                    types.Part.from_bytes(data=video_bytes, mime_type="video/mp4"),
                    types.Part.from_text(text=prompt),
                ]
            )
        ],
        config=types.GenerateContentConfig(
            temperature=0.2,
            max_output_tokens=65536,
        ),
    )
    return response.text

###Scrrenshots
def parse_screenshot_timestamps(analysis_text, chunk_start_sec):
    """Extract [SCREENSHOT:MM:SS] markers from the analysis and return
    list of (absolute_timestamp_sec, description)."""
    pattern = re.compile(r'\[SCREENSHOT:\s*(\d{1,2}):(\d{2})\]\s*(.*?)(?:\n|$)')
    screenshots = []
    for match in pattern.finditer(analysis_text):
        m, s = int(match.group(1)), int(match.group(2))
        relative_sec = m * 60 + s
        absolute_sec = chunk_start_sec + relative_sec
        description = match.group(3).strip().strip('-').strip(':').strip()
        screenshots.append((absolute_sec, relative_sec, description))
    return screenshots


def setup_styles(doc):

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Ensure heading styles exist and look good
    for i, size in [(1, 18), (2, 15), (3, 13)]:
        heading_style = doc.styles[f'Heading {i}']
        heading_style.font.name = 'Calibri'
        heading_style.font.size = Pt(size)
        heading_style.font.color.rgb = RGBColor(0x1A, 0x47, 0x7A)


def markdown_to_docx(doc, markdown_text, screenshots_map=None):

    if screenshots_map is None:
        screenshots_map = {}

    lines = markdown_text.split('\n')
    in_code_block = False
    code_lines = []
    i = 0

    while i < len(lines):
        line = lines[i]

        # Code block toggle
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block — flush
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                p.style = doc.styles['Normal']
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                p_fmt = p.paragraph_format
                from docx.oxml.ns import qn
                shading = p.paragraph_format.element.get_or_add_pPr()
                shd = shading.makeelement(qn('w:shd'), {
                    qn('w:val'): 'clear',
                    qn('w:color'): 'auto',
                    qn('w:fill'): 'F2F2F2',
                })
                shading.append(shd)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Screenshot marker — insert image
        screenshot_match = re.search(r'\[SCREENSHOT:\s*(\d{1,2}):(\d{2})\]', stripped)
        if screenshot_match and screenshots_map:
            m, s = int(screenshot_match.group(1)), int(screenshot_match.group(2))
            ts_key = m * 60 + s
            # Find matching screenshot (check with some tolerance)
            img_path = None
            for abs_ts, info in screenshots_map.items():
                if info['relative_sec'] == ts_key:
                    img_path = info['path']
                    break
            if img_path and os.path.exists(img_path) and os.path.getsize(img_path) > 100:
                # Add description text before image
                desc = re.sub(r'\[SCREENSHOT:\s*\d{1,2}:\d{2}\]\s*', '', stripped)
                if desc:
                    p = doc.add_paragraph()
                    run = p.add_run(desc)
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
                try:
                    doc.add_picture(img_path, width=Inches(6.0))
                    last_paragraph = doc.paragraphs[-1]
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    doc.add_paragraph(f"[Screenshot unavailable at this timestamp]")
                i += 1
                continue

        # Headings
        if stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
            i += 1
            continue
        if stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
            i += 1
            continue
        if stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)
            i += 1
            continue

        # Bullet points
        if stripped.startswith('- ') or stripped.startswith('* '):
            text = stripped[2:]
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, text)
            i += 1
            continue

        # Numbered list
        num_match = re.match(r'^(\d+)\.\s+(.*)', stripped)
        if num_match:
            text = num_match.group(2)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, text)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1


def add_formatted_text(paragraph, text):
    """Add text to a paragraph with basic bold/code inline formatting."""
    # Split on bold and inline code markers
    parts = re.split(r'(\*\*.*?\*\*|`[^`]+`)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
        else:
            paragraph.add_run(part)


def generate_docx(chunk_analyses, all_screenshots, video_path, output_path, total_duration):
    """Generate the final Word document from all chunk analyses."""
    doc = Document()
    setup_styles(doc)

    # ── Title Page ──
    video_name = os.path.splitext(os.path.basename(video_path))[0]
    doc.add_heading(f'KT Meeting Documentation', level=0)
    p = doc.add_paragraph()
    p.add_run(f'Source: ').bold = True
    p.add_run(os.path.basename(video_path))
    p = doc.add_paragraph()
    p.add_run(f'Duration: ').bold = True
    p.add_run(f'{total_duration/60:.0f} minutes')
    p = doc.add_paragraph()
    p.add_run(f'Generated: ').bold = True
    p.add_run(time.strftime('%Y-%m-%d %H:%M'))
    p = doc.add_paragraph()
    p.add_run(f'Chunks Analyzed: ').bold = True
    p.add_run(str(len(chunk_analyses)))

    doc.add_page_break()

    # ── Table of Contents placeholder ──
    doc.add_heading('Table of Contents', level=1)
    doc.add_paragraph('(Update this field in Word: References → Update Table of Contents)')
    doc.add_page_break()

    # ── Per-chunk content ──
    for idx, (analysis, screenshots) in enumerate(zip(chunk_analyses, all_screenshots)):
        # Build screenshot lookup for this chunk
        screenshots_map = {}
        for abs_ts, rel_sec, desc, img_path in screenshots:
            screenshots_map[abs_ts] = {
                'relative_sec': rel_sec,
                'path': img_path,
                'desc': desc
            }

        markdown_to_docx(doc, analysis, screenshots_map)

    doc.save(output_path)
    print(f"\nDocument saved to: {output_path}")


def main():

    VIDEO_PATH      = "/Users/u401396/Code/Utils/gemini/kts/KT  Techops + FTC-20260427_194057-Meeting Recording.mp4"
    TRANSCRIPT_PATH = "/Users/u401396/Code/Utils/gemini/kts/KT _ Techops + FTC.docx"
    OUTPUT_PATH     = None  # None = auto-generates <video_name>_KT_Doc.docx
    MODEL           = "gemini-3.1-pro-preview"
    CHUNK_MINUTES   = 30
    # ════════════════════════════════════════════

    if not os.path.exists(VIDEO_PATH):
        print(f"Error: Video file not found: {VIDEO_PATH}")
        sys.exit(1)

    if OUTPUT_PATH is None:
        video_name = os.path.splitext(os.path.basename(VIDEO_PATH))[0]
        OUTPUT_PATH = os.path.join(os.path.dirname(VIDEO_PATH) or '.', f"{video_name}_KT_Doc.docx")

    chunk_duration_sec = CHUNK_MINUTES * 60

    print("Initializing Gemini client...")
    client = get_gemini_client()

    full_transcript = load_transcript(TRANSCRIPT_PATH)
    if full_transcript:
        print(f"Loaded transcript ({len(full_transcript)} chars)")
    else:
        print("No transcript provided — relying on video understanding only")

    work_dir = tempfile.mkdtemp(prefix="kt_doc_")
    screenshots_dir = os.path.join(work_dir, "screenshots")
    os.makedirs(screenshots_dir, exist_ok=True)

    chunk_paths, total_duration = split_video_into_chunks(
        VIDEO_PATH, chunk_duration_sec, work_dir
    )

    chunk_analyses = []
    all_screenshots = []

    for idx, chunk_path in enumerate(chunk_paths):
        chunk_start_sec = idx * chunk_duration_sec
        chunk_end_sec = min((idx + 1) * chunk_duration_sec, total_duration)

        transcript_segment = get_transcript_segment(
            full_transcript, chunk_start_sec, chunk_end_sec
        )
        previous_summary = chunk_analyses[-1] if chunk_analyses else None
        analysis = analyze_chunk(
            client, MODEL, chunk_path,
            idx, len(chunk_paths), chunk_start_sec, transcript_segment,
            previous_summary=previous_summary,
        )
        chunk_analyses.append(analysis)

        print(f"  Extracting screenshots for chunk {idx + 1}...")
        timestamps = parse_screenshot_timestamps(analysis, chunk_start_sec)
        chunk_screenshots = []

        for abs_ts, rel_sec, desc in timestamps:
            out_path = os.path.join(
                screenshots_dir,
                f"chunk{idx:03d}_{rel_sec:05d}.jpg"
            )
            img_path = extract_screenshot(VIDEO_PATH, abs_ts, out_path)
            if img_path:
                chunk_screenshots.append((abs_ts, rel_sec, desc, img_path))
                print(f"    Screenshot at {abs_ts//60}m{abs_ts%60:02d}s: {desc[:60]}")
            else:
                print(f"    Warning: Could not extract frame at {abs_ts}s")

        all_screenshots.append(chunk_screenshots)
    print("\nGenerating Word document...")
    generate_docx(chunk_analyses, all_screenshots, VIDEO_PATH, OUTPUT_PATH, total_duration)

    print(f"\nDone! Temp files in: {work_dir}")
    print(f"Output: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
 
